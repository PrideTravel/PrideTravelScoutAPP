#!/usr/bin/env python3
"""
Create branded DOCX for Amsterdam Pride 2026 Event Planning Guide
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_heading('Amsterdam Pride 2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format = title.runs[0]
title_format.font.size = Pt(28)
title_format.font.bold = True
title_format.font.color.rgb = RGBColor(255, 87, 34)  # OutAtlas orange

subtitle = doc.add_heading('The Ultimate Event Planning Guide to WorldPride, EuroPride & Pride Amsterdam', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(16)
subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

# Add metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run('Event Dates: July 25 – August 8, 2026  |  Location: Amsterdam, Netherlands')
meta_run.font.size = Pt(10)
meta_run.font.italic = True
meta_run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()  # spacing

# Introduction
intro_heading = doc.add_heading('What Makes Amsterdam Pride 2026 the Most Important LGBTQ+ Celebration in Europe?', level=1)
intro_heading_format = intro_heading.runs[0]
intro_heading_format.font.color.rgb = RGBColor(255, 87, 34)

intro_text = doc.add_paragraph(
    "For the first time in history, Amsterdam will simultaneously host three of the world's most significant Pride celebrations: "
    "Pride Amsterdam, EuroPride, and WorldPride. From July 25 to August 8, 2026, the Dutch capital transforms into a global celebration "
    "of freedom, equality, and queer joy—drawing hundreds of thousands of LGBTQ+ travelers from every corner of the world. "
    "Whether you're attending for the first time or returning for your fifth Pride celebration, this guide covers everything you need "
    "to know to make the most of Amsterdam's legendary two-week festival."
)

# Hotels section
hotels_heading = doc.add_heading('Where Should You Stay During Amsterdam Pride 2026?', level=1)
hotels_heading_format = hotels_heading.runs[0]
hotels_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_paragraph(
    "Amsterdam's accommodations range from luxury waterfront palaces to budget-friendly hostels, all welcoming LGBTQ+ travelers with open arms. "
    "Here's what the market looks like for July 25 – August 8, 2026:"
)

# Create hotel table
table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
headers = ['Hotel Name', 'Star Rating', 'LGBTQ+ Cert.', 'Full Dates', 'Week 1', 'Week 2', 'Best For']
for i, header_text in enumerate(headers):
    hdr_cells[i].text = header_text
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

# Hotel data
hotels_data = [
    ('Anantara Grand Hotel Krasnapolsky', '⭐⭐⭐⭐⭐', 'IGLTA', '$526', '$451', '$526', 'Luxury seekers'),
    ('W Amsterdam', '⭐⭐⭐⭐⭐', 'World Rainbow', '$598', '$518', '$618', 'Modern luxury'),
    ('Pulitzer Amsterdam', '⭐⭐⭐⭐⭐', 'World Rainbow', '$577', '$562', '$577', 'Historic charm'),
    ('Renaissance Amsterdam', '⭐⭐⭐⭐', 'IGLTA', '$331', '$329', '$331', 'Mid-range'),
    ('CitizenM Amsterdam South', '⭐⭐⭐', 'LGBTQ+', '$167', '$163', '$170', 'Budget'),
    ('Volkshotel', '⭐⭐', 'LGBTQ+', '$137', '$135', '$139', 'Backpackers'),
]

for hotel_data in hotels_data:
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(hotel_data):
        row_cells[i].text = cell_text

doc.add_paragraph()

# Add tip
tip = doc.add_paragraph()
tip_run = tip.add_run('Pro tip: ')
tip_run.bold = True
tip.add_run('Week 1 (July 25-31) sees higher demand, but luxury hotels drop rates mid-week. Consider staying July 27-30 for better value at premium properties.')

# Events section
events_heading = doc.add_heading('What Are the Must-See Events During Amsterdam Pride 2026?', level=1)
events_heading_format = events_heading.runs[0]
events_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_paragraph('Amsterdam Pride stretches across 15 days, with major celebrations, performances, conferences, and gatherings scheduled throughout.')

# Events table
events_table = doc.add_table(rows=1, cols=5)
events_table.style = 'Light Grid Accent 1'

event_headers = ['Date', 'Event', 'Time', 'Location', 'Experience']
evt_hdr = events_table.rows[0].cells
for i, header_text in enumerate(event_headers):
    evt_hdr[i].text = header_text
    evt_hdr[i].paragraphs[0].runs[0].font.bold = True

# Event data
events_data = [
    ('Jul 25', 'Pride March', 'All day', 'Amstelveld → Vondelpark', 'Opening celebration'),
    ('Aug 1', 'CANAL PARADE ⭐⭐⭐', '12:00 PM', 'Prinsengracht, Amstel River', '80+ boats, 400K+ spectators'),
    ('Jul 31 - Aug 1', 'World Pride Music Festival ⭐⭐', 'Multi-day', 'Amsterdam venues', 'Electronic music, circuit parties'),
    ('Aug 4-8', 'WorldPride Village', 'All day', 'Museumplein', 'Global hub, performances'),
    ('Aug 8', 'WorldPride Closing', 'Evening', 'Museumplein', 'Grand finale concert'),
]

for event_data in events_data:
    row_cells = events_table.add_row().cells
    for i, cell_text in enumerate(event_data):
        row_cells[i].text = cell_text

doc.add_paragraph()

# Canal Parade section
canal_heading = doc.add_heading('What Is the Canal Parade and Why Is It the World\'s Most Iconic Pride Celebration?', level=1)
canal_heading_format = canal_heading.runs[0]
canal_heading_format.font.color.rgb = RGBColor(255, 87, 34)

canal_sub = doc.add_heading('The History & Significance of Amsterdam\'s Legendary Water Parade', level=2)

canal_intro = doc.add_paragraph(
    "The Canal Parade is what makes Amsterdam Pride truly unique among all Pride celebrations worldwide. Unlike traditional street parades with marching groups and floats, "
    "Amsterdam's Canal Parade takes place entirely on water, with dozens of elaborately decorated boats cruising through the city's UNESCO-protected 17th-century canals. "
    "On August 1, 2026, more than 80 boats will sail the historic Prinsengracht and Amstel rivers, attracting over 400,000 spectators from around the globe."
)

# How it started
doc.add_heading('How It Started', level=3)
doc.add_paragraph(
    "The Canal Parade didn't begin as a protest march—it was intentional celebration. When Amsterdam hosted the Gay Games bid in 1998, organizers envisioned something no other Pride city had ever done: "
    "moving the celebration onto the water. On August 3, 1996, the first Canal Parade launched with just 45 boats and 20,000 spectators. "
    "What emerged was revolutionary: a Pride celebration that transformed the city's most iconic geographic feature (its UNESCO World Heritage canals) into a statement of visibility, joy, and community."
)

# Why it matters
doc.add_heading('Why It Matters', level=3)
doc.add_paragraph(
    "The Canal Parade reflects Amsterdam's unique history. The Netherlands was the first country in the world to legalize same-sex marriage (2001), and Amsterdam specifically has hosted "
    "openly gay establishments since 1927—nearly 100 years of visible queer culture. The city's LGBTQ+ community didn't wait for legal acceptance; it created visibility first. "
    "The Canal Parade embodies that spirit: celebration over petition, joy over struggle, freedom as a lived reality."
)

# The evolution
doc.add_heading('The Evolution', level=3)
doc.add_paragraph(
    "From 1996 onward, the parade grew exponentially. By 2017, it became so culturally significant that organizers renamed 'Amsterdam Gay Pride' to simply 'Pride Amsterdam' to be more inclusive of all LGBTQ+ identities. "
    "In 2019, Pride Amsterdam was officially added to the Inventory of Intangible Heritage in the Netherlands—recognition that this celebration is now part of the country's cultural legacy."
)

doc.add_paragraph(
    "For more information, visit Pride Amsterdam's official website: "
    "https://pride.amsterdam/en/ and explore the Canal Parade details: "
    "https://pride.amsterdam/en/event/canal-parade/"
)

# World Pride Music Festival section
music_heading = doc.add_heading('What Is the World Pride Music Festival (July 31 – August 1)?', level=1)
music_heading_format = music_heading.runs[0]
music_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_paragraph(
    "The World Pride Music Festival is Amsterdam's electronic music and circuit party takeover during Pride week. Running July 31 – August 1, 2026, "
    "the festival features world-class DJs, circuit parties, dance clubs, and themed venues across Amsterdam."
)

doc.add_heading('Festival Highlights:', level=3)
highlights = [
    'Ladz Circuit Party — High-energy men\'s dance party',
    'DAMAGE — Warehouse-style electronic music event',
    'BOPS — Upbeat dance and pop celebration',
    'Rapido — Latino/Caribbean Pride energy',
    'Funhouse — Inclusive dance experience',
    'Bear Necessity — Bear community-focused party',
    'Drag Olympics at Dam Square — Competitive drag performances',
]
for highlight in highlights:
    doc.add_paragraph(highlight, style='List Bullet')

# Legal section
legal_heading = doc.add_heading('Is It Safe to Travel to Amsterdam as an LGBTQ+ Person? What Are My Legal Rights?', level=1)
legal_heading_format = legal_heading.runs[0]
legal_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_heading('Netherlands LGBTQ+ Legal Status (Via Equaldex)', level=2)

doc.add_paragraph(
    'Short answer: Yes, Amsterdam is one of the safest and most welcoming destinations for LGBTQ+ travelers worldwide. '
    'According to Equaldex.com, the Netherlands has the strongest LGBTQ+ legal protections in Europe:'
)

# Legal rights table
legal_table = doc.add_table(rows=1, cols=3)
legal_table.style = 'Light Grid Accent 1'

legal_headers = ['Right/Protection', 'Status', 'Notes']
legal_hdr = legal_table.rows[0].cells
for i, header_text in enumerate(legal_headers):
    legal_hdr[i].text = header_text
    legal_hdr[i].paragraphs[0].runs[0].font.bold = True

# Legal data
legal_data = [
    ('Same-Sex Marriage', '✅ Legal since 2001', 'First country globally'),
    ('Same-Sex Adoption', '✅ Legal', 'Joint, stepchild, fertility access'),
    ('Anti-Discrimination', '✅ Comprehensive', 'Employment, housing, public access'),
    ('Transgender Recognition', '✅ Legal', 'Gender identity protections'),
    ('Public Support', '✅ 97% approval', '2019 Eurobarometer data'),
]

for legal_item in legal_data:
    row_cells = legal_table.add_row().cells
    for i, cell_text in enumerate(legal_item):
        row_cells[i].text = cell_text

doc.add_paragraph()
doc.add_paragraph(
    'For detailed information, visit Equaldex: LGBT Rights in Netherlands at '
    'https://www.equaldex.com/region/netherlands'
)

# Budget section
budget_heading = doc.add_heading('What\'s the Budget for Amsterdam Pride 2026?', level=1)
budget_heading_format = budget_heading.runs[0]
budget_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_paragraph('Here\'s a realistic cost breakdown for a 7-night / 8-day stay (July 25–August 2):')

# Budget table
budget_table = doc.add_table(rows=1, cols=4)
budget_table.style = 'Light Grid Accent 1'

budget_headers = ['Category', 'Budget', 'Moderate', 'Luxury']
budget_hdr = budget_table.rows[0].cells
for i, header_text in enumerate(budget_headers):
    budget_hdr[i].text = header_text
    budget_hdr[i].paragraphs[0].runs[0].font.bold = True

budget_data = [
    ('Hotel (7 nights)', '$945', '$2,303', '$4,039'),
    ('Meals (3x daily)', '$560', '$1,120', '$1,680'),
    ('Pride Events', '$150', '$300', '$500+'),
    ('Attractions/Tours', '$100', '$250', '$400+'),
    ('Transportation', '$50', '$75', '$100'),
    ('Nightlife/Bars', '$200', '$400', '$600+'),
    ('Shopping/Misc', '$100', '$250', '$500+'),
    ('TOTAL (USD)', '$2,105', '$4,698', '$7,819+'),
]

for budget_item in budget_data:
    row_cells = budget_table.add_row().cells
    for i, cell_text in enumerate(budget_item):
        row_cells[i].text = cell_text

doc.add_paragraph()

# Money-saving tips
tip_head = doc.add_paragraph()
tip_head_run = tip_head.add_run('Money-saving tips:')
tip_head_run.bold = True

tips = [
    'Book accommodations 3–4 months in advance for better rates',
    'Many Pride events and Canal Parade viewpoints are FREE',
    'Bike rentals ($10–15/day) are cheaper than taxis',
    'Restaurant lunches are 30–40% cheaper than dinners',
    'Visit during week 1 (July 25–31) for slightly lower hotel rates',
]
for tip in tips:
    doc.add_paragraph(tip, style='List Bullet')

# Conclusion
doc.add_page_break()

conclusion_heading = doc.add_heading('Final Thoughts: Why Amsterdam Pride 2026 Is Unmissable', level=1)
conclusion_heading_format = conclusion_heading.runs[0]
conclusion_heading_format.font.color.rgb = RGBColor(255, 87, 34)

doc.add_paragraph(
    "The convergence of Pride Amsterdam, EuroPride, and WorldPride in 2026 is unprecedented. For 15 days, the Dutch capital becomes the world's epicenter of LGBTQ+ celebration, activism, and community. "
    "The Canal Parade on August 1 isn't just a parade—it's a statement that visibility, joy, and freedom are not granted by law alone; they're lived, celebrated, and defended by communities year after year.\n\n"
    "Whether this is your first Pride or your fifteenth, Amsterdam in July–August 2026 offers something no other destination can: a celebration rooted in nearly 100 years of queer history, supported by a nation that made equality law before it was mainstream, "
    "and shared with hundreds of thousands of LGBTQ+ people from every corner of the globe.\n\n"
    "Book your hotel, pack your Pride gear, and get ready for the most important LGBTQ+ event in the world."
)

final_note = doc.add_paragraph()
final_note_run = final_note.add_run('See you in Amsterdam. 🏳️‍🌈')
final_note_run.font.size = Pt(14)
final_note_run.font.bold = True
final_note_run.font.color.rgb = RGBColor(255, 87, 34)

# Resources
doc.add_page_break()

resources_heading = doc.add_heading('Resources & Links', level=1)
resources_heading_format = resources_heading.runs[0]
resources_heading_format.font.color.rgb = RGBColor(255, 87, 34)

resources = [
    ('Pride Amsterdam Official Website', 'https://pride.amsterdam/en/'),
    ('WorldPride Amsterdam 2026', 'https://pride.amsterdam/en/worldpride/'),
    ('EuroPride 2026 Amsterdam', 'https://epoa.eu/europride/europride-2026-amsterdam/'),
    ('World Pride Music Festival', 'https://worldpridefest.com/'),
    ('Equaldex: Netherlands LGBT Rights', 'https://www.equaldex.com/region/netherlands'),
    ('Homomonument Amsterdam', 'https://www.homomonument.nl/'),
    ('COC Amsterdam Community Center', 'https://www.coc.nl/'),
    ('IHLIA LGBT Heritage Museum', 'https://www.ilgia.org/'),
    ('Zandvoort Gay Beach', 'https://zandvoort.biz/'),
    ('Boat Parade Tickets', 'https://www.botenparade.nl/'),
]

for resource_name, resource_url in resources:
    p = doc.add_paragraph()
    run = p.add_run(f'{resource_name}: ')
    run.bold = True
    p.add_run(resource_url)

# Save document
doc.save('Amsterdam_Pride_2026_Event_Planning_Guide.docx')
print('✅ Document created: Amsterdam_Pride_2026_Event_Planning_Guide.docx')
